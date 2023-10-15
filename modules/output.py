import os
import io
import random
import zipfile
from classes.activity import activity
from classes.placeholder import placeholder
from pywebio.input import *
from pywebio.output import *
from pywebio.pin import *
from time import time
from modules.base64_coder import *
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from PIL import Image
import docx
from docxcompose.composer import Composer

def _replace(placeholder: str, paragraph, text: str):
    if placeholder in paragraph.text:
        paragraph.text = paragraph.text.replace(placeholder, text)

def output(act: activity):
    while True:
        action = actions(label=f"导出活动[{act.name}]的信息", buttons=[
            {'label': "下载JSON存档", 'value': "to_json"},
            {'label': "生成DOCX材料", 'value': "to_docx"},
            {'label': "导出所有电子发票", 'value': "output_invoice"},
            {'label': "返回上一级", 'value': "back"}
        ], help_text=f"详细信息：{act.info()}")
        value_map = {
            "to_json": to_json,
            "to_docx": to_docx,
            "output_invoice": output_invoice,
            "back": None
        }
        if value_map[action] is None:
            break
        else:
            value_map[action](act)

# 导出所有电子发票
def output_invoice(act: activity):
    # 格式选择：所有发票打包成zip，或者放进一个pdf单独导出
    # format = select("请选择导出格式：", options=["zip", "pdf"], required=True, help_text="zip格式可以一次性下载所有发票，所有发票的图片将被打包成一个zip文件；pdf格式将所有发票放进一个pdf文件中")
    format = "zip"
    # 根据选择进行处理
    if format == "zip":
        # 生成zip文件
        try:
            _flag = False
            time_stamp = int(time())
            random_num = random.randint(1000,9999)
            zip_name = f"./temp_invoice_{time_stamp}_{random_num}.zip"
            # 创建zip文件
            with zipfile.ZipFile(zip_name, "w") as zip:
                for item in act.items:
                    if item.invoice is not None:
                        # TODO: 上传发票时，不一定全是JPG(将发票写入内存时，建议全部转成JPG格式)
                        _flag = True
                        content = from_base64(item.invoice)
                        zip.writestr(f"{item.name}.jpg", content)
            if _flag == False:
                toast("没有电子发票可以下载")
                return
            # 提供下载
            put_file(f"{act.name}_{int(time())}.zip", content=open(zip_name, "rb").read())
        except Exception as e:
            print(e)
            toast("下载失败，请联系管理员")
        finally:
            if zip_name not in os.listdir():
                return
            # 删除zip文件
            os.remove(zip_name)
    else:
        # TODO: 生成pdf文件
        toast("暂不支持pdf格式导出")

# 下载json存档
def to_json(act: activity):
    serialized = act.to_json()
    # 提供文件下载
    put_file(f"{act.name}_{int(time())}.json",
             content=serialized.encode("utf-8"))
    # 退出
    return None

# 将活动输出为docx
def to_docx(act: activity):
    if len(act.items) == 0:
        toast("活动物资为空，无法生成申请表")
        return

    try:
        # 计算act预算，按照用途
        budget = {}
        for item in act.items:
            if item.use not in budget.keys():
                budget[item.use] = 0.0
            budget[item.use] += item.total
        # 按照预算排序
        budget = sorted(budget.items(), key=lambda item: item[1], reverse=True)
        # 如果预算超过了6个，只取前5个，剩下的归类为其他
        if len(budget) > 6:
            other_sum = sum([item[1] for item in budget[6:]])
            budget = budget[:6]
            budget.append(("其他", other_sum))
        # 生成保存文件
        time_stamp = int(time())
        random_num = random.randint(1000,9999)
        template_name = "./templates/template_applicant_form.docx"
        form_name = f"./temp_form_{time_stamp}_{random_num}.docx"
        # fill the template file
        doc = Document(template_name)
        for table in doc.tables:
            for row in table.rows:
                for cell in set(row.cells):
                    for paragraph in cell.paragraphs:
                        _replace(placeholder.act_name, paragraph, act.name)
                        _replace(placeholder.act_time, paragraph, act.time)
                        _replace(placeholder.act_place, paragraph, act.place)
                        _replace(placeholder.act_plan, paragraph, "见后附页" if act.plan else "暂无")
                        _replace(placeholder.act_date, paragraph, act.time.replace("-", "年", 1).replace("-", "月", 1))
                        _replace(placeholder.dir_name, paragraph, act.director.name)
                        _replace(placeholder.dir_phone, paragraph, act.director.phone)
                        _replace(placeholder.dir_mail, paragraph, act.director.email)
                        _replace(placeholder.itm_cnt, paragraph, str(len(budget)))
                        _replace(placeholder.itm_sum, paragraph, str(sum([item[1] for item in budget])))
                        for i in range(1, 7):
                            if i <= len(budget):
                                _replace(f"{placeholder.itm_grp}{i}", paragraph, budget[i-1][0])
                                _replace(f"{placeholder.itm_bgt}{i}", paragraph, str(budget[i-1][1]))
                            else:
                                _replace(f"{placeholder.itm_grp}{i}", paragraph, "")
                                _replace(f"{placeholder.itm_bgt}{i}", paragraph, "")

        # 策划书
        if act.plan is not None:
            doc.add_page_break()
            plan = Document(io.BytesIO(act.plan))
            plan.add_page_break()
        else:
            plan = doc

        # 物资清单
        title = plan.add_paragraph()
        title.alignment = 1  # 设置为居中对齐（1表示居中）
        run = title.add_run("活动物资清单")
        run.font.size = Pt(18)

        table = plan.add_table(rows=1, cols=7)
        # set the table header
        headers = ['组别', '名称', '数量', '单价', '总价', '用途', '备注']
        for idx, header in enumerate(headers):
            table.cell(0, idx).text = header

        for it in act.items:
            cells = table.add_row().cells
            values = [it.group, it.name, it.count, it.price, it.total, it.use, it.note]
            for idx, value in enumerate(values):
                cells[idx].text = str(value)

        # 组内合计
        rowIndex = len(table.rows)
        table.add_row()
        table.cell(rowIndex, 0).merge(table.cell(rowIndex, len(headers)-1))
        table.cell(rowIndex, 0).text = "组内合计"

        # 各组总计
        for group in act.groups:
            rowIndex = len(table.rows)
            table.add_row()
            table.cell(rowIndex, 0).merge(table.cell(rowIndex, 1))
            table.cell(rowIndex, 2).merge(table.cell(rowIndex, len(headers)-1))
            table.cell(rowIndex, 0).text = group
            table.cell(rowIndex, 2).text = str(sum(it.total if it.total else 0 for it in act.items if it.group == group))

        # 合计
        rowIndex = len(table.rows)
        table.add_row()
        table.cell(rowIndex, 0).merge(table.cell(rowIndex, 5))
        table.cell(rowIndex, 0).text = "合计"
        table.cell(rowIndex, 6).text = str(sum(it.total if it.total else 0 for it in act.items))

        # 增加边框
        for row in table.rows:
            for cell in row.cells:
                # 获取单元格的tcPr对象
                tcPr = cell._tc.get_or_add_tcPr()

                # 创建边框元素
                borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                    r'<w:top w:val="single" w:color="000000" w:sz="8"/>'
                                    r'<w:bottom w:val="single" w:color="000000" w:sz="8"/>'
                                    r'<w:left w:val="single" w:color="000000" w:sz="8"/>'
                                    r'<w:right w:val="single" w:color="000000" w:sz="8"/>'
                                    r'</w:tcBorders>')

                # 将边框样式应用到单元格
                tcPr.append(borders)

        # 组别合并
        index = [-1]
        for i in range(len(act.items) - 1):
            if act.items[i].group != act.items[i + 1].group:
                index.append(i)
        index.append(len(act.items) - 1)
        for i in range(len(index) - 1):
            start, end = index[i] + 1, index[i + 1]
            table.cell(start + 1, 0).merge(table.cell(end + 1, 0))
            table.cell(start + 1, 0).text = act.items[start].group



        # 发票
        for item in act.items:
            if item.invoice is not None:
                plan.add_page_break()
                title = plan.add_paragraph()
                run = title.add_run(f"{item.name}的发票，用途：{item.use}")
                run.font.size = Pt(12)
                # invoice 是 JPG 格式的 bytes 转成的 base64 字符串
                invoice = from_base64(item.invoice)
                invoice = io.BytesIO(invoice)
                invoice_img = Image.open(invoice)
                rotated_image_stream = io.BytesIO()
                rotated_image = invoice_img.rotate(-90, expand=True)
                rotated_image.save(rotated_image_stream, format='PNG')
                rotated_image_stream.seek(0)
                plan.add_picture(rotated_image_stream, width=Inches(5.5), height=Inches(9))
            else:
                # 如果没有发票，就添加template_paper_invoice.docx
                template_name = "./templates/template_paper_invoice.docx"
                template = Document(template_name)
                for element in template.element.body:
                    plan.element.body.append(element)

        plan.add_page_break()

        composer = Composer(doc)
        composer.append(plan)

        # 新闻稿
        if act.news is not None:
            news = Document(io.BytesIO(act.news))
            composer.append(news)

        # save the file
        composer.save(form_name)

        # provide download
        put_file(f"{act.name}_{int(time())}.docx", content=open(form_name, "rb").read())
    except Exception as e:
        print(e)
        toast("下载失败，请联系管理员")
    finally:
        if form_name not in os.listdir():
            return
        # remove the file
        os.remove(form_name)