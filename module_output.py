from objects import *
from pywebio.input import *
from pywebio.output import *
from pywebio.pin import *
from time import time
from base64_coder import *
from docx import Document
from dataclasses import dataclass
import os
import random

@dataclass
class placeholder:
    act_name: str = "PLACEHOLDER_ACTIVITY_NAME"
    act_time: str = "PLACEHOLDER_ACTIVITY_TIME"
    act_date: str = "PLACEHOLDER_ACTIVITY_DATE"
    act_place: str = "PLACEHOLDER_ACTIVITY_PLACE"
    act_plan: str = "PLACEHOLDER_SCHEME"
    dir_name: str = "PLACEHOLDER_DIRECTOR_NAME"
    dir_phone: str = "PLACEHOLDER_DIRECTOR_PHONE"
    dir_mail: str = "PLACEHOLDER_DIRECTOR_MAIL"
    itm_grp: str = "PLACEHOLDER_ITEM_GROUP"
    itm_bgt: str = "PLACEHOLDER_ITEM_BUDGET"
    itm_cnt: str = "PLACEHOLDER_ITEM_COUNT"
    itm_sum: str = "PLACEHOLDER_ITEM_SUMUP"


def fill_applicant_form(act: activity):
    # 计算act预算，按照用途
    budget = {}
    for item in act.item:
        if item.use not in budget.keys():
            budget[item.use] = 0.0
        budget[item.use] += item.total
    # 按照预算排序
    budget = sorted(budget.items(), key=lambda item: item[1], reverse=True)
    # 如果预算超过了6个，只取前5个，剩下的归类为其他
    if len(budget) > 6:
        other_sum = sum([item[1] for item in budget[6:]])
        budget = budget[:6]
        budget.append(("其他", other_sum))
    # 生成保存文件
    time_stamp = int(time())
    random_num = random.randint(1000,9999)
    template_name = "./template_applicant_form.docx"
    form_name = f"./temp_form_{time_stamp}_{random_num}.docx"
    # fill the template file
    doc = Document(template_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder.act_name in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.act_name, act.name)
                    if placeholder.act_time in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.act_time, act.time)
                    if placeholder.act_place in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.act_place, act.place)
                    if placeholder.act_plan in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.act_plan, "见后附页" if act.plan else "暂无")
                    if placeholder.act_date in paragraph.text:
                        # 将yyyy-mm-dd转换为yyyy年mm月dd日
                        date_copy = act.date.copy()
                        date_copy = date_copy.replace("-", "年", 1)
                        date_copy = date_copy.replace("-", "月", 1)
                        paragraph.text = paragraph.text.replace(placeholder.act_date, date_copy)
                    if placeholder.dir_name in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.dir_name, act.director.name)
                    if placeholder.dir_phone in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.dir_phone, act.director.phone)
                    if placeholder.dir_mail in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.dir_mail, act.director.email)
                    for i in range(1, 7):
                        if f"{placeholder.itm_grp}{i}" in paragraph.text:
                            if i <= len(budget):
                                paragraph.text = paragraph.text.replace(f"{placeholder.itm_grp}_{i}", budget[i-1][0])
                            else:
                                paragraph.text = paragraph.text.replace(f"{placeholder.itm_grp}_{i}", "")
                        if f"{placeholder.itm_bgt}{i}" in paragraph.text:
                            if i <= len(budget):
                                paragraph.text = paragraph.text.replace(f"{placeholder.itm_bgt}_{i}", str(budget[i-1][1]))
                            else:
                                paragraph.text = paragraph.text.replace(f"{placeholder.itm_bgt}_{i}", "")
                    if placeholder.itm_cnt in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.itm_cnt, str(len(budget)))
                    if placeholder.itm_sum in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder.itm_sum, str(sum([item[1] for item in budget])))
    # save the file
    doc.save(form_name)
    # provide download
    put_file(f"{act.name}_{int(time())}.docx", content=open(form_name, "rb").read())